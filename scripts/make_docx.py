#!/usr/bin/env python3
"""生成 Word 版:build/issue.json → radar/word/音乐科技雷达_第NNN期_YYYY-MM-DD.docx

排版对照样例《音乐科技雷达_第001期_2026-07-23.docx》:
- 页面 Letter,边距左右 0.945" 上下 0.866"
- 字体:西文 Times New Roman,中文宋体
- 报名 26pt 粗体报头红 C1301C 居中;副信息 9pt 灰 666666 居中
- 板块标题「■ ××」14pt 粗体红;条目标题 12pt 粗体油墨 1A1A1A
- meta/链接 8.5pt 灰;正文 10.5pt 油墨
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Emu, Pt, RGBColor
except ImportError:
    print("[docx] 错误:未安装 python-docx。请执行:\n"
          "  python3 -m venv radar/.venv && radar/.venv/bin/pip install -r radar/requirements.txt\n"
          "然后用 radar/.venv/bin/python 运行本脚本。", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
ISSUE = ROOT / "build" / "issue.json"
OUT_DIR = ROOT / "word"

RED = RGBColor(0xC1, 0x30, 0x1C)
INK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)

WEEKDAYS = "一二三四五六日"


def style_run(run, size, bold=False, color=INK):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    # 中文字体:宋体
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def para(doc, text, size, bold=False, color=INK, center=False, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), size, bold, color)
    return p


def add_item(doc, idx_title, it):
    """一条收录条目:标题 / meta(+英文原题)/ detail 三段 / 原文链接。"""
    para(doc, idx_title, 12, bold=True, after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    meta = f"〔{it.get('tag', '')}〕 {it.get('src', '')} · {it.get('date', '')}"
    run = p.add_run(meta)
    if it.get("orig"):
        run.add_break()
        run.add_text(it["orig"])
    style_run(run, 8.5, color=GRAY)
    for seg in (it.get("detail") or it.get("why") or "").split("\n"):
        if seg.strip():
            para(doc, seg.strip(), 10.5, after=4)
    para(doc, f"原文链接:{it.get('url', '')}", 8.5, color=GRAY, after=4)


def main():
    ap = argparse.ArgumentParser(description="生成当期 Word 版")
    ap.add_argument("--vol", type=int, help="期号(默认读 issue.json)")
    ap.add_argument("--date", help="出版日期(默认读 issue.json)")
    args = ap.parse_args()

    issue = json.loads(ISSUE.read_text(encoding="utf-8"))
    vol = args.vol or issue.get("vol")
    pub = args.date or issue.get("pub")
    if not vol or not pub:
        print("[docx] 错误:期号/日期未知(请先运行 build_page.py 或传 --vol/--date)",
              file=sys.stderr)
        sys.exit(1)
    y, m, d = (int(x) for x in pub.split("-"))
    from datetime import date as date_cls
    weekday = f"星期{WEEKDAYS[date_cls(y, m, d).weekday()]}"

    items = issue["items"]
    ideas = issue["ideas"]
    hot_paper = next((it for it in items if it.get("hot") == "本报头条"), None)
    hot_ind = next((it for it in items if it.get("hot") == "行业要闻"), None)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Emu(864235)
    sec.top_margin = sec.bottom_margin = Emu(791845)

    # ---- 报头 ----
    para(doc, "音乐科技雷达", 26, bold=True, color=RED, center=True, after=2)
    para(doc, f"第 {vol:03d} 期 · {y}年{m}月{d}日 {weekday} · "
              f"收录 {len(items)} 条 · 选题 {len(ideas)} 个", 9, color=GRAY, center=True, after=2)
    para(doc, "数字音乐智能处理 · 学术前沿 × 开源工具 × 行业动态",
         9, color=GRAY, center=True, after=8)

    # ---- 本报头条 ----
    para(doc, "■ 本报头条", 14, bold=True, color=RED, after=8)
    for hot in (hot_paper, hot_ind):
        if hot:
            add_item(doc, f"★{hot['hot']}★ {hot['title']}", hot)

    # ---- 三个板块 ----
    for sec_key, sec_name in (("papers", "学术前沿"), ("oss", "开源工具"), ("industry", "行业动态")):
        rows = [it for it in items if it["sec"] == sec_key and not it.get("hot")]
        if not rows:
            continue
        para(doc, f"■ {sec_name}", 14, bold=True, color=RED, after=8)
        for i, it in enumerate(rows, 1):
            add_item(doc, f"{i}. {it['title']}", it)

    # ---- 选题灵感 ----
    if ideas:
        para(doc, "■ 选题灵感(基于近期各期积累 · 按难度从低到高)",
             14, bold=True, color=RED, after=8)
        for idea in ideas:
            para(doc, f"{idea.get('level', '')} · {idea.get('kind', '')}|{idea.get('title', '')}",
                 12, bold=True, after=4)
            para(doc, idea.get("body", ""), 10.5, after=4)

    # ---- 报尾 ----
    para(doc, "数据来源:arXiv(cs.SD / eess.AS / cs.HC / cs.MM)· Hacker News · GitHub"
              "|按专业相关度筛选,宁缺毋滥", 8.5, color=GRAY, after=2)
    para(doc, f"音乐科技雷达 · GENERATED {pub} · 每日 07:00 出版", 8.5, color=GRAY, after=2)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"音乐科技雷达_第{vol:03d}期_{pub}.docx"
    doc.save(out)
    # 重新打开确认无损坏
    chk = Document(out)
    print(f"[docx] {out}({len(chk.paragraphs)} 段,{out.stat().st_size} 字节)")


if __name__ == "__main__":
    main()
